import asyncio
from pathlib import Path
from docx import Document
from pypdf import PdfReader
from services.translator import translator

def read_document(path: str) -> list[str]:
    suffix = Path(path).suffix.lower()
    if suffix == ".txt":
        return [Path(path).read_text(encoding="utf-8", errors="replace")]
    if suffix == ".docx":
        doc = Document(path)
        return [p.text for p in doc.paragraphs if p.text.strip()]
    if suffix == ".pdf":
        return [(page.extract_text() or "") for page in PdfReader(path).pages]
    raise ValueError("Faqat PDF, DOCX va TXT fayllar qabul qilinadi.")

async def translate_document(input_path: str, output_dir: str, source: str, target: str) -> str:
    parts = await asyncio.to_thread(read_document, input_path)
    if not any(x.strip() for x in parts):
        raise ValueError("Hujjatdan matn topilmadi. Skaner PDF bo‘lsa, sahifalarni rasm qilib yuboring.")
    translated = []
    for part in parts:
        if part.strip():
            chunks = [part[i:i+3500] for i in range(0, len(part), 3500)]
            translated.append("\n".join([await translator.translate(c, source, target) for c in chunks]))
    source_path = Path(input_path)
    if source_path.suffix.lower() == ".txt":
        out = Path(output_dir) / f"{source_path.stem}_{target}.txt"
        await asyncio.to_thread(out.write_text, "\n\n".join(translated), "utf-8")
    else:
        out = Path(output_dir) / f"{source_path.stem}_{target}.docx"
        doc = Document()
        doc.add_heading("Tarjima", 0)
        for text in translated:
            doc.add_paragraph(text)
        await asyncio.to_thread(doc.save, out)
    return str(out)

